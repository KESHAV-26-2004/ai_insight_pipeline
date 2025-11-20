# report_generator.py
import os
import json
import re
from pathlib import Path
import pandas as pd

# For DOCX
from docx import Document
from docx.shared import Inches, Pt

# For PDF
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.units import inch

# -----------------------------
# Load JSON Helper
# -----------------------------
def load_json(path):
    if not os.path.exists(path):
        return {}
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)

def create_sentiment_pie_chart(df, sentiment_col, out_path):
    """
    Create a simple pie chart for sentiment distribution and save to out_path.
    """
    import matplotlib.pyplot as plt

    try:
        counts = df[sentiment_col].value_counts(dropna=True).to_dict()
    except Exception:
        return False

    if not counts:
        return False

    labels_map = {
        0: "Negative",
        1: "Neutral",
        2: "Positive",
        "negative": "Negative",
        "neutral": "Neutral",
        "positive": "Positive"
    }

    values = []
    names = []

    for k, v in counts.items():
        names.append(labels_map.get(k, str(k)))
        values.append(v)

    plt.figure(figsize=(5, 5))
    plt.pie(values, labels=names, autopct="%1.1f%%", startangle=140)
    plt.title("Sentiment Distribution")
    plt.tight_layout()
    # ensure parent dir exists
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    plt.savefig(out_path, dpi=150)
    plt.close()
    return True

# -----------------------------
# Build Markdown (for reference, optional)
# -----------------------------
def write_markdown(md_path, dataset_dir, title, filename, profile, sentiment_counts, sentiment_pie, relations, relation_sentences, summary, recommendations):
    md = []

    md.append(f"# {title}\n")
    md.append(f"### Dataset: **{filename}**\n")

    md.append("## 1. Dataset Overview\n")
    md.append(f"- **Rows:** {profile.get('n_rows')}")
    md.append(f"- **Columns:** {len(profile.get('columns', {}))}\n")

    md.append("### Column Details\n")
    for col, info in profile.get("columns", {}).items():
        sample = info["sample_values"][0] if info.get("sample_values") else "-"
        md.append(f"- **{col}** ({info['dtype']}) — sample: `{sample}`")

    md.append("\n## 2. Sentiment Summary\n")
    for k, v in sentiment_counts.items():
        md.append(f"- **{k}**: {v}")

    if sentiment_pie:
        md.append(f"\n![Sentiment Pie]({Path(sentiment_pie).as_posix()})")

    md.append("\n## 3. Key Drivers & Correlations\n")
    for r, s in zip(relations, relation_sentences):
        md.append(f"- **{s.get('long','')}**")
        if "plot" in r:
            normalized = r["plot"].replace("\\", "/")
            full_plot = Path(dataset_dir) / normalized
            md.append(f"  \n![Relation Plot]({full_plot.as_posix()})")

    md.append("\n## 4. Summary\n")
    md.append(summary)

    md.append("\n## 5. Recommendations\n")
    for i, rec in enumerate(recommendations, start=1):
        md.append(f"{i}. {rec}")

    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(md))

# -----------------------------
# DOCX Generation
# -----------------------------
def generate_docx(docx_path, dataset_dir, title, filename, profile, sentiment_counts, sentiment_pie, relations, relation_sentences, summary, recommendations):
    doc = Document()

    # Dataset filename (centered)
    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run(filename)
    run.font.size = Pt(14)
    run.bold = True

    # Title (H1)
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run(title)
    h1_run.font.size = Pt(24)

    doc.add_heading("1. Dataset Overview", level=2)
    doc.add_paragraph(f"Rows: {profile.get('n_rows')}")
    doc.add_paragraph(f"Columns: {len(profile.get('columns', {}))}")

    doc.add_heading("Column Details", level=3)
    for col, info in profile.get("columns", {}).items():
        sample = info["sample_values"][0] if info.get("sample_values") else "-"
        doc.add_paragraph(f"{col} ({info['dtype']}): {sample}")

    doc.add_heading("2. Sentiment Summary", level=2)
    for k, v in sentiment_counts.items():
        doc.add_paragraph(f"{k}: {v}")

    if sentiment_pie and os.path.exists(sentiment_pie):
        try:
            doc.add_picture(str(sentiment_pie), width=Inches(5))
        except Exception:
            # fail silently if docx cannot embed image
            pass

    doc.add_heading("3. Key Drivers & Correlations", level=2)
    for r, s in zip(relations, relation_sentences):
        # use long sentence as requested
        long_text = s.get("long", s.get("short", ""))
        doc.add_paragraph(long_text)

        if "plot" in r:
            normalized = r["plot"].replace("\\", "/")
            img_path = Path(dataset_dir) / normalized
            if img_path.exists():
                try:
                    doc.add_picture(str(img_path), width=Inches(5))
                except Exception:
                    pass

    doc.add_heading("4. Summary", level=2)
    for line in str(summary).split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_heading("5. Recommendations", level=2)
    for i, rec in enumerate(recommendations, start=1):
        doc.add_paragraph(f"{i}. {rec}")

    # ensure parent dir exists
    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    doc.save(docx_path)

# -----------------------------
# PDF Generation (Clean layout)
# -----------------------------
def generate_pdf(pdf_path, dataset_dir, title, filename, profile, sentiment_counts, sentiment_pie, relations, relation_sentences, summary, recommendations):
    styles = getSampleStyleSheet()
    story = []

    # Custom styles
    title_style = ParagraphStyle(
        name="TitleStyle",
        parent=styles["Heading1"],
        alignment=TA_CENTER,
        fontSize=22,
        spaceAfter=20
    )

    filename_style = ParagraphStyle(
        name="FilenameStyle",
        parent=styles["Normal"],
        alignment=TA_CENTER,
        fontSize=14,
        spaceAfter=10,
        leading=16
    )

    h2 = styles["Heading2"]
    h3 = styles["Heading3"]
    normal = styles["BodyText"]

    # Filename
    story.append(Paragraph(filename, filename_style))
    story.append(Paragraph(title, title_style))

    # Section 1
    story.append(Paragraph("1. Dataset Overview", h2))
    story.append(Paragraph(f"Rows: {profile.get('n_rows')}", normal))
    story.append(Paragraph(f"Columns: {len(profile.get('columns', {}))}", normal))
    story.append(Spacer(1, 12))

    # Column details
    story.append(Paragraph("Column Details", h3))
    for col, info in profile.get("columns", {}).items():
        sample = info["sample_values"][0] if info.get("sample_values") else "-"
        story.append(Paragraph(f"<b>{col}</b> ({info['dtype']}): {sample}", normal))
    story.append(Spacer(1, 14))

    # Sentiment Summary
    story.append(Paragraph("2. Sentiment Summary", h2))
    for k, v in sentiment_counts.items():
        story.append(Paragraph(f"{k}: {v}", normal))

    if sentiment_pie and os.path.exists(sentiment_pie):
        story.append(Spacer(1, 10))
        try:
            story.append(Image(str(sentiment_pie), width=4*inch, height=4*inch))
        except Exception:
            pass
    story.append(Spacer(1, 16))

    # Relations
    story.append(Paragraph("3. Key Drivers & Correlations", h2))
    for r, s in zip(relations, relation_sentences):
        # long sentence
        story.append(Paragraph(s.get("long", s.get("short", "")), normal))
        if "plot" in r:
            normalized = r["plot"].replace("\\", "/")
            img_path = Path(dataset_dir) / normalized
            if img_path.exists():
                story.append(Spacer(1, 10))
                try:
                    story.append(Image(str(img_path), width=5.5*inch, height=3*inch))
                except Exception:
                    pass
                story.append(Spacer(1, 16))

    # Summary
    story.append(Paragraph("4. Summary", h2))
    for line in str(summary).split("\n"):
        if line.strip():
            story.append(Paragraph(line.strip(), normal))
    story.append(Spacer(1, 16))

    # Recommendations
    story.append(Paragraph("5. Recommendations", h2))
    for i, rec in enumerate(recommendations, start=1):
        story.append(Paragraph(f"{i}. {rec}", normal))
    story.append(Spacer(1, 20))

    # ensure parent dir exists
    os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
    doc = SimpleDocTemplate(pdf_path, pagesize=letter)
    doc.build(story)

# -----------------------------
# MAIN REPORT GENERATOR
# -----------------------------
def _detect_sentiment_column_from_meta_and_df(sentiment_meta, df):
    """
    Try to detect a proper sentiment column to use (numeric preferred).
    """
    # Check sentiment_meta existing_sentiment structure
    try:
        existing = sentiment_meta.get("existing_sentiment", [])
        if isinstance(existing, list) and existing:
            first = existing[0]
            if isinstance(first, dict) and first.get("numeric_version") in df.columns:
                return first.get("numeric_version")
            if isinstance(first, dict) and first.get("column") in df.columns:
                return first.get("column")
        # sometimes sentiment_meta might store sentiment_columns list
        cand = sentiment_meta.get("sentiment_columns") or sentiment_meta.get("sentiment_column")
        if cand:
            if isinstance(cand, list) and len(cand) > 0 and cand[0] in df.columns:
                return cand[0]
            if isinstance(cand, str) and cand in df.columns:
                return cand
    except Exception:
        pass

    # fallback: detect any column in df that endswith _sentiment_num first then _sentiment
    for c in df.columns:
        if c.endswith("_sentiment_num"):
            return c
    for c in df.columns:
        if c.endswith("_sentiment"):
            return c
    # last resort: common name
    for name in ("content_sentiment_num", "content_sentiment", "sentiment_num", "sentiment"):
        if name in df.columns:
            return name
    return None

def _sanitize_summary(raw_summary):
    if not raw_summary:
        return ""
    s = str(raw_summary)
    # Remove variants of FINAL SUMMARY header (case-insensitive, with optional asterisks or hashes)
    s = re.sub(r"(?im)^\s*(\*{0,2}\s*final summary\s*\*{0,2}\s*[:\-]*\s*)", "", s)
    s = re.sub(r"(?im)^\s*(#{1,3}\s*final summary\s*)", "", s)
    # Remove any remaining literal "**FINAL SUMMARY**"
    s = s.replace("**FINAL SUMMARY**", "")
    s = s.replace("FINAL SUMMARY", "")
    # Trim whitespace
    return s.strip()

def generate_full_report(dataset_dir):
    dataset_dir = str(dataset_dir).rstrip("/\\")
    base = Path(dataset_dir).name
    filename = f"{base}.csv"

    # Load components
    title = load_json(f"{dataset_dir}/title/{base}_title.json").get("generated_title", base)
    profile = load_json(f"{dataset_dir}/profiles/{base}_profile.json")
    sentiment_meta = load_json(f"{dataset_dir}/enriched/{base}_sentiment_meta.json")
    relations = load_json(f"{dataset_dir}/relations/{base}_relations.json") or []
    relation_sentences = load_json(f"{dataset_dir}/relations/{base}_relations_sentences.json") or []
    gemini = load_json(f"{dataset_dir}/gemini/{base}_gemini_output.json")

    # Features CSV
    features_csv = f"{dataset_dir}/features/{base}_features.csv"
    if not os.path.exists(features_csv):
        raise FileNotFoundError(f"Features CSV not found: {features_csv}")

    df = pd.read_csv(features_csv, low_memory=False)

    # Sentiment column detection
    sentiment_col = _detect_sentiment_column_from_meta_and_df(sentiment_meta, df)
    sentiment_counts = {}
    sentiment_pie = None

    # Ensure plots folder exists
    plots_dir = Path(dataset_dir) / "plots"
    os.makedirs(plots_dir, exist_ok=True)

    if sentiment_col and sentiment_col in df.columns:
        try:
            sentiment_counts = df[sentiment_col].value_counts().to_dict()
        except Exception:
            sentiment_counts = {}
        # create pie chart file path and generate pie
        sentiment_pie_path = str(plots_dir / "sentiment_pie.png")
        try:
            ok = create_sentiment_pie_chart(df, sentiment_col, sentiment_pie_path)
            if ok and os.path.exists(sentiment_pie_path):
                sentiment_pie = sentiment_pie_path
        except Exception:
            sentiment_pie = None
    else:
        # No sentiment detected: keep counts empty
        sentiment_counts = {}

    # Summary cleaning (remove FINAL SUMMARY variants)
    raw_summary = gemini.get("summary", "") if isinstance(gemini, dict) else ""
    clean_summary = _sanitize_summary(raw_summary)

    # Output folder
    report_dir = f"{dataset_dir}/report"
    os.makedirs(report_dir, exist_ok=True)

    md_path = f"{report_dir}/{base}_report.md"
    pdf_path = f"{report_dir}/{base}_report.pdf"
    docx_path = f"{report_dir}/{base}_report.docx"

    # Write MD
    write_markdown(md_path, dataset_dir, title, filename, profile, sentiment_counts, sentiment_pie, relations, relation_sentences, clean_summary, gemini.get("recommendations", []))

    # DOCX
    generate_docx(docx_path, dataset_dir, title, filename, profile, sentiment_counts, sentiment_pie, relations, relation_sentences, clean_summary, gemini.get("recommendations", []))

    # PDF
    generate_pdf(pdf_path, dataset_dir, title, filename, profile, sentiment_counts, sentiment_pie, relations, relation_sentences, clean_summary, gemini.get("recommendations", []))

    return {"md": md_path, "pdf": pdf_path, "docx": docx_path}

# -----------------------------
# CLI
# -----------------------------
if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--dataset_folder", required=True)
    args = parser.parse_args()

    out = generate_full_report(args.dataset_folder)
    print("Report generated:", out)
